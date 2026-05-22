#!/usr/bin/env python3
"""
Build tumbler-paper-gtac2026.docx from the GTAC submission template
by python-docx.  Strategy:

  1. Load the template so theme/styles/header/footer come along intact.
  2. Wipe all body content (paragraphs + tables) without touching the XML
     parts that hold the header/footer (those live in separate parts).
  3. Re-emit our content using the template's named styles so the
     document keeps the GTAC look-and-feel.
"""
import sys, os
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pathlib import Path as _Path
REPO = _Path(__file__).resolve().parents[1]


TEMPLATE = str(REPO / 'docs/templates/GTAC2025-Submission-Template-Final.docx')
OUT      = str(REPO / 'docs/paper/tumbler-paper-gtac2026.docx')
FIG_DIR  = REPO / 'docs/paper/figures'

doc = Document(TEMPLATE)


def clear_body(doc):
    """Remove every block-level element from the body except the final
    sectPr (page setup of section 0).  We rebuild section structure
    explicitly via add_section() below."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

clear_body(doc)


def set_columns(section, num, space=540):
    """Force a section to N columns by editing <w:cols> on its sectPr."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space))


def disable_numbering(paragraph):
    """Explicitly suppress style-inherited auto-numbering on this paragraph
    by setting <w:numPr><w:numId w:val='0'/></w:numPr> (numId=0 = no list)."""
    pPr = paragraph._p.get_or_add_pPr()
    for n in pPr.findall(qn('w:numPr')):
        pPr.remove(n)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '0')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)


def switch_to_columns(num):
    """Insert a CONTINUOUS section break.  The previous section freezes
    here; a new section starts that takes effect from the NEXT paragraph.
    Sets the new section to <num> columns."""
    new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_columns(new_section, num)
    return new_section


# ---------- helpers ----------
def add_para(text='', style='Normal', align=None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        if bold:
            r.bold = True
    if align is not None:
        p.alignment = align
    return p


def add_picture(path, width_in=3.0, caption=None):
    """Default figure width tuned for the GTAC 2-column body
    (each column is ~3.25" wide).  Slide-shaped (16:9) source images at
    3.0" wide are about 1.7" tall, so a figure + caption occupies
    roughly 18% of one body page and fits comfortably in a single
    column without overflowing the gutter."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(caption, style='Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_h1(title):
    """Heading 1 — template auto-numbers as '1.', '2.', ... so we write
    only the title text."""
    return add_para(title, style='Heading 1')


def add_h2(title):
    return add_para(title, style='Heading 2')


def add_unnumbered_h1(title):
    """Heading 1 styling but with numbering suppressed (for ABSTRACT,
    References, etc. that should not appear in the auto-numbered list)."""
    p = add_para(title, style='Heading 1')
    disable_numbering(p)
    return p


def add_bullets(items, style='List Paragraph'):
    for it in items:
        add_para(it, style=style)


def _set_cell_border(cell, *, sz='6', color='000000'):
    """Apply an explicit single black border to all 4 sides of a cell.
    sz is in eighths-of-a-point ('6' = 0.75 pt).  Works around Word
    table-styles whose borders don't render in LibreOffice."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_table_borders(table, *, sz='6', color='000000'):
    """Apply explicit single black borders on the table-level <w:tblBorders>
    so the lines render in Word AND LibreOffice regardless of the
    selected table style."""
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._element.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_table(headers, rows, style='Table Grid', col_widths_in=None):
    """Add a table sized to fill the full text-area width of the current
    section, with explicit black single borders on all cells.

    NOTE: we don't trust template table styles to render borders in
    LibreOffice — instead we always emit explicit <w:tblBorders> +
    per-cell <w:tcBorders>.  Style 'Table Grid' is used as a sensible
    base, but the explicit borders win over the style."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = 'Table Grid'  # neutral base, our borders override
    except KeyError:
        pass

    # Pin table width to 100% of the containing text area.
    tblPr = t._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); t._element.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')   # 5000 = 100.00% in fiftieths
    tblPr.append(tblW)

    # Explicit black borders for the whole table.
    _set_table_borders(t)

    # Explicit per-column widths (if given) — render Word/LibreOffice
    # more deterministically than autofit alone.
    if col_widths_in is not None:
        from docx.shared import Inches as _In
        for j, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[j].width = _In(w)

    def _tighten(paragraph, bold=False, header=False):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(9 if not header else 10)
            if bold:
                run.font.bold = True

    # Also paint borders on every cell explicitly (defensive: some
    # renderers honour cell-level borders ahead of table-level ones).
    for row in t.rows:
        for cell in row.cells:
            _set_cell_border(cell)

    # header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        _tighten(p, bold=True, header=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[1 + i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.add_run(val)
            _tighten(p)
    return t


# ============================================================
#                         CONTENT
# ============================================================
#
# Section layout we build:
#   Section 0: 1-column  -- title + authors + abstract + keywords
#   Section 1: 2-column  -- body
#   (Section breaks around wide tables temporarily switch back to 1-column.)
#
# The template's "Heading 1" style auto-numbers as "1.", "2.", ... via a
# multilevel list that ALSO drives Heading 7.  Without explicit suppression
# Heading 7 (ABSTRACT) renders "1.1.1.1.1.1.1" and pushes the Heading 1
# counter past 1.  We use disable_numbering() on every heading that should
# not appear in the section auto-number sequence.

# Tell section 0 to be a single column (the kept sectPr is currently 2-col).
set_columns(doc.sections[0], 1)


# ------------- Title block (single-column) -------------
add_para('TUMBLER: A Layered Survival Firewall for the AMD ROCm Stack '
         'under Production AI Workloads',
         style='Title')

add_para('Chun-Hung Wang*, Clement Lin, Jeremy Liao', style='Subtitle')

add_para('AMD', style='Body Text')
add_para('Chun-Hung.Wang@amd.com, Clement.Lin@amd.com, Jeremy.Liao@amd.com',
         style='Body Text')
add_para('* Corresponding author', style='Body Text')

# ------------- Abstract (single-column) -------------
p_abs = add_para('ABSTRACT', style='Heading 7')
disable_numbering(p_abs)
add_para(
    'Production AI serving stacks share AMD GPUs across many concurrent '
    'tenants, run for hours or days continuously, and treat transient '
    'hardware and driver faults as a routine operational cost. The default '
    'ROCm runtime, designed for HPC single-tenant correctness, calls '
    'std::abort() on a VM fault, a wedged SDMA ring, or a parked HSA signal '
    '— turning a localized fault into a process-wide outage that takes every '
    'co-tenant down with it. We argue that this is the right behaviour for '
    'the workload ROCm was originally designed for, and the wrong behaviour '
    'for production AI. We present TUMBLER — a Three-tier Unified '
    'Multi-tenant Bounded-wait Layered Escape-hatch Runtime — an opt-in '
    'survival firewall layered across the three lines of the ROCm stack '
    '(the CLR HIP host runtime, the ROCr HSA runtime, and the amdgpu '
    'kernel driver) that bounds the wait time at every known abort or '
    'unbounded-park site and returns a graceful failure to the calling '
    'request instead of aborting the process. Tumbler defaults to off '
    '(byte-identical to stock) so it is safe to land upstream. We have '
    'submitted Tumbler as five independent '
    'pull requests against ROCm/rocm-systems and ROCm/amdgpu, and validated '
    'the integrated stack on a multi-tenant MI300X production serving '
    'workload that previously wedged within single-digit minutes; with '
    'Tumbler enabled, the same workload remained stable for over 11 hours.',
    style='Normal')

add_para(
    'Keywords: AI serving, ROCm, AMD GPU, fault tolerance, runtime '
    'survivability, HIP, HSA, amdgpu kernel driver, MI300X.',
    style='Quote')

# --- switch to 2-column layout for the body ---
switch_to_columns(2)

# ------------- 1. Introduction -------------
add_section_h1('Introduction')

add_para(
    'Picture yourself driving on the highway. The vehicle infotainment '
    'touchscreen — rendered by AMD amdgpu — hits a transient GPU fault '
    'mid-frame. What you want from that system is unambiguous: drop the bad '
    'frame, log the fault, and keep the dashboard, the windshield HUD, and '
    'the surround-view cameras alive. What you do not want is for the GPU '
    'driver to crash the entire in-vehicle infotainment process and freeze '
    'every safety-critical display you depend on. The argument is intuitive '
    'when the failure surface is a windshield. It is the same argument when '
    'the failure surface is an AI inference fleet serving thousands of '
    'concurrent tenants on a node of MI300X GPUs.', style='Normal')

add_picture(FIG_DIR / 'fig-02-car-cockpit-frozen.png',
            caption='Figure 1. The vehicle cockpit argument. The central '
                    'infotainment touchscreen (amdgpu) has hit a GPU fault, '
                    'but the dashboard cluster and windshield HUD continue '
                    'rendering normally. Drop the bad pipeline, keep the '
                    'safety surfaces alive.')

add_para(
    'Modern AI serving stresses the AMD GPU runtime in ways it was not '
    'originally designed for: a single inference process hosts hundreds of '
    'concurrent tenant requests sharing the same GPUs, SDMA engines, HSA '
    'signals, and KFD pinned-memory windows. Long-running tensor ops, '
    'multi-tenant sharing, RDMA-Write pipelines, and eviction during '
    'inference all produce transient HW/driver faults whose blast radius '
    'in stock ROCm is the whole process — every co-tenant inside dies even '
    'though only one request was affected. This is a workload-shape '
    'mismatch rather than a bug: ROCm was built for HPC single-tenant '
    'correctness, where a VM fault means data corruption and aborting is '
    'exactly the right response. Production AI inverts the assumption — '
    'the cost of killing the process is several orders of magnitude '
    'larger than the cost of dropping one request — and needs an opt-in '
    'escape hatch that does not change default behaviour for legitimate '
    'single-tenant HPC users.', style='Normal')

add_para(
    'We present Tumbler, a layered opt-in survival firewall for the ROCm '
    'stack. Tumbler installs bounded-wait knobs at every known abort site '
    'and every known unbounded-park site across the three layers that '
    'matter — the CLR host runtime, the ROCr HSA runtime, and the amdgpu '
    'kernel driver — and converts those sites from "abort the process" or '
    '"block the caller forever" into "fail this one call gracefully and '
    'let the higher-level scheduler decide what to do". Every knob '
    'defaults to off, so a system that does not set any Tumbler knob '
    'behaves byte-identically to stock ROCm.', style='Normal')

add_para('Contributions:', style='Normal', bold=True)
add_bullets([
    'We characterize the workload-shape mismatch between HPC '
    'single-tenant correctness and production AI serving, and identify '
    'the abort/park sites it exposes across the ROCm stack.',
    'We design Tumbler, a layered opt-in survival firewall spanning '
    'CLR, ROCr, amdkcl, and amdgpu, comprising 10 self-contained '
    'environment-variable and modparam knobs.',
    'We measure the firewall on a real multi-tenant MI300X serving '
    'workload and show that a stack that previously wedged in single-'
    'digit minutes can be made stable for over 11 hours with Tumbler '
    'knobs enabled.',
    'We submit the entire design upstream as five independent pull '
    'requests against ROCm/rocm-systems and ROCm/amdgpu, and publish '
    'an integrated meta-repository so the combination can be rebuilt '
    'deterministically.',
])

# ------------- 2. Background -------------
add_section_h1('Background: The ROCm Stack Today')

add_para(
    'AMD ROCm is the open-source GPU compute stack for AMD Instinct and '
    'Radeon hardware. From the application down to the device, three '
    'layers matter for survivability analysis: CLR (the HIP host runtime, '
    'a.k.a. rocclr) manages streams and waits on HSA signals; ROCr (the '
    'HSA runtime) owns queues, signals, SDMA copies, and the user-mode '
    'side of fault delivery; and amdgpu (the upstream Linux kernel '
    'driver) manages VRAM, GTT, and KFD, with an amdkcl helper layer '
    'providing DKMS shims for non-mainline kernel APIs. A request issued '
    'by an AI serving framework travels down through these three layers '
    'into the GPU and back. At every layer there is a wait point where '
    'the caller can be parked indefinitely if the lower layer fails to '
    'make progress, and historically an abort site that escalates a fault '
    'into a process-wide termination.', style='Normal')

# ------------- 3. Why ROCm has no firewall -------------
add_section_h1('Why ROCm Has No Built-in Survival Firewall')

add_para(
    'The HSA specification, the ancestor of ROCr, was designed with the '
    'assumption that one logical user process owns the GPU. Under that '
    'assumption, a hardware VM fault, a queue exception, or a memory '
    'corruption signal from the device is by definition unrecoverable — '
    'the calling process has by hypothesis already lost data integrity, '
    'and the only safe action is to abort.', style='Normal')

add_para(
    'ROCr implements this faithfully. The three abort sites in the runtime '
    '— the memory critical-error handler, the HW exception handler, and '
    'the VM fault handler — all call std::abort() unconditionally. CLR\'s '
    'WaitForSignal() polling loop has a 4-second per-call timeout in its '
    'call to Hsa::signal_wait_scacquire(), but the outer while loop has no '
    'wall-clock cap, so a permanently wedged signal parks the caller '
    'thread forever. SDMA writers spin on os::YieldThread() until they '
    'can publish an address. The amdgpu kernel driver waits on '
    'dma_fence_wait_any_timeout() with MAX_SCHEDULE_TIMEOUT inside '
    'kcl_drm_suballoc_new() and on a global mutex inside the GTT window '
    'path. Every one of these is correct behaviour when one process owns '
    'one GPU.', style='Normal')

add_para(
    'There is one existing knob that approaches the survivability problem '
    'from a different angle: HSA_SIGNAL_WAIT_ABORT_TIMEOUT. When set, the '
    'HSA runtime aborts the process when a signal wait exceeds the '
    'configured deadline. This is the right tool to backstop runaway '
    'kernels in a batch HPC job. It is the wrong tool for AI serving: it '
    'still aborts the process, just on a timer instead of on a fault. A '
    'serving stack that aborts at the deadline loses every co-tenant '
    'request in flight — even though those requests are still healthy.',
    style='Normal')

# ------------- 4. Failure modes -------------
add_section_h1('Failure Modes Under AI Serving Load')

add_para(
    'We observed five distinct failure modes during multi-tenant production '
    'serving on customer MI300X nodes: SDMA ring stall (writers spin in '
    'AcquireWriteAddress() on os::YieldThread() indefinitely); per-request '
    'VM fault (ROCr\'s handler calls std::abort() and every co-tenant '
    'dies with it); HSA signal wait parked (CLR\'s WaitForSignal() outer '
    'loop has no wall-clock cap); GTT window starvation (every '
    'VRAM-touching ioctl parks on the global mman.gtt_window_lock mutex '
    'held by one in-flight SDMA copy); and KFD orphan pin (a pinned BO\'s '
    'owning process exits without unpinning, holding VRAM open and '
    'starving subsequent allocations). Stack signatures and root causes '
    'vary, but the pattern is uniform: one request hits an unrecoverable '
    'condition, the stack escalates it to a process-wide abort or an '
    'indefinite park, and every co-tenant request is collateral damage.',
    style='Normal')

# ------------- 5. Tumbler design -------------
add_section_h1('Tumbler Design — A Layered Survival Firewall')

add_picture(FIG_DIR / 'fig-04-tumbler-firewall-layered.png',
            caption='Figure 2. Tumbler installs a bounded-wait firewall at '
                    'each layer of the stack. The fault is stopped at the '
                    'layer closest to it, the affected request fails '
                    'gracefully, and unrelated co-tenant requests continue '
                    'serving.')

add_para(
    'Tumbler treats each ROCm layer as a firewall plane. At every plane '
    'we install (a) a bounded-wait conversion of the corresponding '
    'unbounded-park site, and (b) where the plane is an abort site, an '
    'opt-in non-abort branch that surfaces the failure as a return code '
    'instead of as a process-wide std::abort(). Two design rules are '
    'fixed across the whole stack: the default is off (byte-identical '
    'to stock), and the failure mode at the deadline is graceful — the '
    'calling code receives a well-defined error, never an abort.',
    style='Normal')

add_para(
    'The protections decompose cleanly along the three boundary planes '
    'of the runtime — CLR, ROCr, and amdgpu — each catching a specific '
    'class of failure as it tries to propagate up the stack. Table 1 '
    'summarises what each layer guards against and how it does so. We '
    'leave individual knob names and tuning details to the upstream '
    'pull requests cited in Section 7 and to the meta-repository, so '
    'that this paper focuses on the protection model rather than on '
    'configuration mechanics.', style='Normal')

switch_to_columns(1)  # full-width band for the protection table

add_table(
    headers=['Layer', 'Guards against', 'How it protects'],
    rows=[
        ['CLR\n(HIP host runtime)',
         'A host thread parked forever on a wedged HSA completion '
         'signal; the serving scheduler cannot reroute or fail the '
         'affected request.',
         'Bounded outer-loop wait on WaitForSignal(). At the deadline '
         'the signal is force-completed and the caller returns, '
         'letting the upper layer decide whether to retry or fail '
         'the one request.'],
        ['ROCr\n(HSA runtime)',
         'std::abort() at the three runtime.cpp abort sites '
         '(memory error, non-ECC HW exception, VM fault) — one fault '
         'takes down the whole process and every co-tenant request '
         'with it. Plus SDMA writers spinning forever in '
         'AcquireWriteAddress(), and InterruptSignal waits parked.',
         'Opt-in non-abort branches at every abort site (ECC stays '
         'abort-fatal for data integrity). Bounded caps on the '
         'signal-wait inner loop and on the SDMA yield loop. '
         'Failure surfaces to the caller as an HSA_STATUS error, '
         'never as a process abort.'],
        ['amdgpu\n(Linux kernel driver)',
         'A wedged SDMA ring holding the global GTT-window mutex '
         'and starving every VRAM-touching ioctl on the device. '
         'Pinned BOs whose owner exited without unpinning, '
         'permanently holding VRAM. Unbounded dma_fence_wait '
         'inside suballoc.',
         'Bounded timeouts at every fence and lock site (suballoc, '
         'GTT window, KFD free / unpin). KFD pin-reaper background '
         'subsystem harvests orphan pins on a deadline. Failure '
         'surfaces to user space as -ETIME, never as a hung ioctl.'],
    ],
    style='Steve\'s Technical Table',
    col_widths_in=[1.5, 2.5, 2.8])

add_para("Table 1. Tumbler's three layers of defense.", style='Caption',
         align=WD_ALIGN_PARAGRAPH.CENTER)

switch_to_columns(2)  # resume 2-column body

# ------------- 6. Evaluation -------------
add_section_h1('Evaluation')

add_para(
    'We validated the integrated Tumbler stack on a customer MI300X serving '
    'box running a representative production AI inference workload. The '
    'baseline (stock ROCm) wedged within single-digit minutes in our '
    'reproducer; with Tumbler knobs enabled at the gold values, the same '
    'workload ran continuously for over 11 hours.', style='Normal')

switch_to_columns(1)  # full-width band for the waveform chart
add_picture(FIG_DIR / 'fig-05-latency-waveform.png', width_in=6.6,
            caption='Figure 3. Serving-workload waveform on the same '
                    'AI inference reproducer.  Both runs start at the '
                    'same baseline (\u2248 75 % GPU util, \u2248 50 ms '
                    'p99 latency).  At t = 8 min a transient fault hits.  '
                    'Stock ROCm aborts: util and latency go away (red '
                    'band).  Tumbler shows a brief bounded latency '
                    'spike at the ROCR_SIGNAL_WAIT_MAX_MS = 500 ms cap, '
                    'then the workload recovers; subsequent hiccups '
                    'behave the same way.')
switch_to_columns(2)  # resume 2-column body

add_para(
    'We further isolated the ROCr non-abort signal-wait cap '
    '(ROCR_SIGNAL_WAIT_MAX_MS) under a microbenchmark that holds the '
    'signal wedged. The measured time-to-return scales linearly with the '
    'configured cap (Table 2), confirming that the deadline mechanism is '
    'sharp and predictable.', style='Normal')

switch_to_columns(1)  # full-width band for the signal-wait scaling table

add_table(
    headers=['ROCR_SIGNAL_WAIT_MAX_MS', 'Measured wall-clock to return'],
    rows=[
        ['500', '515 ms'],
        ['200', '213 ms'],
        ['100', '114 ms'],
    ],
    style='Steve\'s Technical Table',
    col_widths_in=[3.0, 3.0])

add_para('Table 2. Sub-second linear scaling of the ROCr signal-wait cap.',
         style='Caption', align=WD_ALIGN_PARAGRAPH.CENTER)

switch_to_columns(2)  # resume 2-column body

add_para(
    'Beyond raw stability, the firewall isolates faults per-request rather '
    'than per-process: a fault hitting one tenant\'s request on one GPU '
    'fails that single request gracefully, while every co-tenant on the '
    'same MI300X node continues serving uninterrupted.', style='Normal')

# ------------- 7. Upstream status -------------
add_section_h1('Upstream Status')

add_para(
    'The entire Tumbler design has been submitted upstream as five '
    'independent pull requests so that each protection plane can be '
    'reviewed and merged on its own merit: ROCm/rocm-systems #6328 '
    '(CLR) and #6329 (ROCr), and ROCm/amdgpu #214 (amdkcl), #215 '
    '(amdgpu GTT), and #216 (amdgpu KFD pin reaper). No PR depends on '
    'any other; each can be enabled in isolation.', style='Normal')

add_para(
    'The integrated combination of all five is published as the meta-'
    'repository github.com/AFDEAPAC/Tumbler with two submodules pinned '
    'at integration tips (chun-wan/rocm-systems @ tumbler-integrated '
    'for the CLR + ROCr changes and chun-wan/amdgpu @ tumbler-integrated '
    'for the three amdgpu changes), so the exact stack measured in '
    'Section 6 can be rebuilt deterministically.', style='Normal')

# ------------- 8. Related work -------------
add_section_h1('Related Work')

add_para(
    'HSA_SIGNAL_WAIT_ABORT_TIMEOUT is the closest existing upstream knob; '
    'it bounds the wait but escalates at the deadline to a process-wide '
    'abort, which is the behaviour Tumbler deliberately avoids for AI '
    'serving. The two are complementary: set ABORT_TIMEOUT strictly larger '
    'than ROCR_SIGNAL_WAIT_MAX_MS to keep a hard backstop on top of the '
    'graceful cap. Windows GPU TDR resets the GPU on long-running kernels '
    'but is opaque to the user-mode runtime and offers no per-request '
    'isolation inside a single host process. Linux IOMMU recovery isolates '
    'faulting devices via page-fault routing one layer below the GPU '
    'runtime and does not address user-mode abort sites in CLR or ROCr. '
    'NVIDIA MPS partitions a single GPU across multiple processes for '
    'inter-process isolation; Tumbler is intra-process isolation for the '
    'case where dozens of tenants share one Python serving process by '
    'design.', style='Normal')

# ------------- 9. Conclusion -------------
add_section_h1('Conclusion and Future Work')

add_para(
    'Tumbler is small. The aggregate diff is on the order of a few hundred '
    'lines of added code spread across three repos, every knob defaults '
    'to off, and every knob is independently reviewable. The argument is '
    'not that the implementation is novel — the argument is that the '
    'underlying assumption (one host process, one GPU, fault = data '
    'corruption) is no longer the workload that matters most, and that '
    'the runtime needs a small, opt-in set of escape hatches to be '
    'usable as the substrate for production AI.', style='Normal')

add_para(
    'Future work centres on integrating Tumbler with the AI workload '
    'itself rather than leaving knob values as static environment '
    'variables: auto-tuning each deadline from observed workload '
    'signatures (tail-latency budget, batch dwell time, in-flight '
    'count), reinforcement-learning-driven survival policies, and '
    'feedback channels that surface graceful-fail events back into the '
    'inference framework so the higher-level scheduler can adjust load.',
    style='Normal')

# ------------- References (unnumbered) -------------
add_unnumbered_h1('References')
add_para(
    '[1] ROCm/rocm-systems PR #6328 — rocclr: optional bounded HSA signal '
    'wait via HIP_MAX_SIGNAL_WAIT. https://github.com/ROCm/rocm-systems/'
    'pull/6328', style='Normal')
add_para(
    '[2] ROCm/rocm-systems PR #6329 — rocr: non-abort service-survival '
    'envs. https://github.com/ROCm/rocm-systems/pull/6329', style='Normal')
add_para(
    '[3] ROCm/amdgpu PR #214 — drm/amdkcl: bounded kcl_drm_suballoc_new '
    'wait via suballoc_timeout_ms. https://github.com/ROCm/amdgpu/pull/214',
    style='Normal')
add_para(
    '[4] ROCm/amdgpu PR #215 — drm/amdgpu: bounded mman.gtt_window_lock '
    'acquisition via gtt_lock_timeout_ms. https://github.com/ROCm/amdgpu/'
    'pull/215', style='Normal')
add_para(
    '[5] ROCm/amdgpu PR #216 — drm/amdgpu: KFD RDMA-pin bounded-wait + '
    'orphan reaper subsystem. https://github.com/ROCm/amdgpu/pull/216',
    style='Normal')
add_para(
    '[6] AFDEAPAC/Tumbler — integrated meta-repository. '
    'https://github.com/AFDEAPAC/Tumbler', style='Normal')
add_para(
    '[7] ROCm contributing guidelines. https://github.com/ROCm/ROCm/blob/'
    'develop/CONTRIBUTING.md#pull-requests', style='Normal')
add_para(
    '[8] HSA Foundation. HSA Runtime Programmer\'s Reference Manual.',
    style='Normal')

# ============================================================
#                          SAVE
# ============================================================
# --- core properties (Tumbler) -----------------------------------------
# Override the GTAC template's original docProps/core.xml so reviewers
# see the right author/subject/title in File -> Properties.
cp = doc.core_properties
cp.title = 'TUMBLER: A Three-tier Bounded-Wait Survival Stack for ROCm under Transient HW/Driver Faults'
cp.subject = 'AMD ROCm CLR + ROCr + amdgpu fault-tolerant runtime/driver stack'
cp.author = 'Chun-Hung Wang'
cp.last_modified_by = 'Chun-Hung Wang'
cp.keywords = 'ROCm; HIP; HSA; amdgpu; bounded-wait; survival; multi-tenant; AI serving; MI300X'
cp.comments = 'GTAC 2026 submission. Generated by scripts/build_paper.py from the GTAC2025 template.'
# -----------------------------------------------------------------------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'wrote: {OUT}')
print(f'  size: {os.path.getsize(OUT)} bytes')


# --- scrub docProps/app.xml leftover template fields (Tumbler) -------
import zipfile, shutil, re as _re, os as _os, tempfile as _tempfile
def _scrub_app_xml(path):
    tmp_fd, tmp_path = _tempfile.mkstemp(suffix='.xlsx')
    _os.close(tmp_fd)
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'docProps/app.xml':
                txt = data.decode('utf-8', errors='ignore')
                txt = _re.sub(r'<Manager>[^<]*</Manager>', '<Manager></Manager>', txt)
                txt = _re.sub(r'<Company>\s*</Company>', '<Company>Advanced Micro Devices</Company>', txt)
                data = txt.encode('utf-8')
            zout.writestr(item, data)
    shutil.move(tmp_path, path)
    _os.chmod(path, 0o644)

_scrub_app_xml(OUT)
print('scrubbed app.xml placeholders in', OUT)
# ---------------------------------------------------------------------
