#!/usr/bin/env python3
from docx import Document
d = Document('/tmp/tumbler-build/tumbler-paper-gtac2025.docx')
n_paras = sum(1 for _ in d.paragraphs)
n_tables = len(d.tables)
n_imgs = 0
for p in d.paragraphs:
    for r in p.runs:
        if r._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            n_imgs += 1
# Approximate word count
words = 0
for p in d.paragraphs:
    words += len(p.text.split())
print(f"paragraphs: {n_paras}")
print(f"tables: {n_tables}")
print(f"images:  {n_imgs}")
print(f"words:   {words}  (approx page count @ 500 wpm: {words/500:.1f})")
print()
print("=== first 25 non-empty paragraphs ===")
shown = 0
for p in d.paragraphs:
    if p.text.strip():
        print(f"  [{p.style.name}] {p.text[:90]}")
        shown += 1
        if shown >= 25:
            break
